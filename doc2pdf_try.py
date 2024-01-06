import os
import argparse
from docx import Document
from docx2pdf import convert

# Parse arguments from the command line
parser = argparse.ArgumentParser(description='Replace specific placeholders and convert to PDF')
parser.add_argument('--docx_file', help='Path to the .docx file')
parser.add_argument('--value1', help='Replacement for Value_1')
parser.add_argument('--value2', help='Replacement for Value_2')
# Add more arguments as needed
args = parser.parse_args()

# Open the .docx file
doc = Document(args.docx_file)

# Replace specific placeholders with the given replacements
print("doc.paragraphs",doc.paragraphs)
for para in doc.paragraphs:
    print("para",para)
    if 'Value1' in para.text and args.value1:
        para.text = para.text.replace('Value1=', f"Value1={args.value1}")
    if 'Value2' in para.text and args.value2:
        para.text = para.text.replace('Value2=', f"Value2={args.value2}")
    # Add more replacement rules as needed

# Save the modified .docx
modified_docx_path = 'modified_test.docx'
doc.save(modified_docx_path)

# Convert the modified .docx to PDF
pdf_path = modified_docx_path.replace('.docx', '.pdf')
convert(modified_docx_path, pdf_path)


os.remove('modified_test.docx')
print(f'Successfully converted {args.docx_file} to {pdf_path}')
